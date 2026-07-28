from pathlib import Path

from docx import Document
from docx.shared import Pt
from pydantic import BaseModel

from app.schemas import Evidence, TailoringResult
from app.schemas.models import SupportStatus


class DocxExportSummary(BaseModel):
    output_path: str
    claim_count: int
    evidence_count: int


def export_tailored_docx(
    result: TailoringResult,
    output_path: str | Path,
    *,
    title: str = "Tailored Resume Claims",
) -> DocxExportSummary:
    invalid_claims = [
        claim for claim in result.claims if claim.support_status != SupportStatus.SUPPORTED
    ]
    if invalid_claims:
        raise ValueError("DOCX export only accepts claims with supported status.")
    if not result.claims:
        raise ValueError("DOCX export requires at least one approved claim.")

    evidence_by_id: dict[str, Evidence] = {
        evidence.source_id: evidence
        for mapping in result.evidence_map
        for evidence in mapping.evidence
    }
    referenced_ids = {evidence_id for claim in result.claims for evidence_id in claim.evidence_ids}
    missing_ids = sorted(referenced_ids - evidence_by_id.keys())
    if missing_ids:
        raise ValueError(f"Claims reference missing evidence IDs: {', '.join(missing_ids)}")

    destination = Path(output_path)
    destination.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Aptos"
    normal_style.font.size = Pt(10.5)
    document.add_heading(title, level=0)
    document.add_paragraph(
        "Only human-approved, evidence-supported claims are included in this document."
    )

    if result.application_summary:
        document.add_heading("Application Summary", level=1)
        document.add_paragraph(result.application_summary)

    if result.cover_letter:
        document.add_heading("Cover Letter", level=1)
        for block in result.cover_letter.split("\n\n"):
            document.add_paragraph(block)
    document.add_heading("Tailored Claims", level=1)
    for claim in result.claims:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(claim.text)
        source_run = paragraph.add_run(f"\nEvidence: {', '.join(claim.evidence_ids)}")
        source_run.italic = True
        source_run.font.size = Pt(8.5)

    document.add_page_break()  # type: ignore[no-untyped-call]
    document.add_heading("Evidence Audit", level=1)
    for evidence_id in sorted(referenced_ids):
        evidence = evidence_by_id[evidence_id]
        document.add_heading(evidence_id, level=2)
        document.add_paragraph(f"Source: {evidence.source_path}")
        document.add_paragraph(evidence.excerpt)

    document.save(str(destination))
    return DocxExportSummary(
        output_path=str(destination),
        claim_count=len(result.claims),
        evidence_count=len(referenced_ids),
    )

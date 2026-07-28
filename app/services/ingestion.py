import hashlib
from collections.abc import Callable
from io import BytesIO
from pathlib import Path

from app.schemas import EvidenceDocument


class UnsupportedDocumentError(ValueError):
    """Raised when a source file has no registered loader."""


def _source_id(path: Path) -> str:
    digest = hashlib.sha256(str(path.resolve()).encode()).hexdigest()[:12]
    return f"{path.stem.lower().replace(' ', '-')}-{digest}"


def _uploaded_source_id(filename: str, content: bytes) -> str:
    stem = Path(filename).stem.lower().replace(" ", "-")
    digest = hashlib.sha256(content).hexdigest()[:12]
    return f"{stem}-{digest}"


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as error:
        raise RuntimeError("Install document extras with: uv sync --extra documents") from error
    return "\n\n".join(page.extract_text() or "" for page in PdfReader(path).pages)


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as error:
        raise RuntimeError("Install document extras with: uv sync --extra documents") from error
    document = Document(str(path))
    blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    return "\n\n".join(blocks)


LOADERS: dict[str, Callable[[Path], str]] = {
    ".md": _read_text,
    ".markdown": _read_text,
    ".txt": _read_text,
    ".pdf": _read_pdf,
    ".docx": _read_docx,
}


def load_document(path: str | Path) -> EvidenceDocument:
    source_path = Path(path)
    if not source_path.is_file():
        raise FileNotFoundError(f"Document not found: {source_path}")
    loader = LOADERS.get(source_path.suffix.casefold())
    if loader is None:
        supported = ", ".join(sorted(LOADERS))
        raise UnsupportedDocumentError(
            f"Unsupported document type '{source_path.suffix}'. Supported: {supported}"
        )
    content = loader(source_path).strip()
    if not content:
        raise ValueError(f"No extractable text found in: {source_path}")
    return EvidenceDocument(
        source_id=_source_id(source_path),
        source_path=str(source_path),
        content=content,
    )


def load_documents(paths: list[str | Path]) -> list[EvidenceDocument]:
    return [load_document(path) for path in paths]


def load_uploaded_document(filename: str, content: bytes) -> EvidenceDocument:
    """Parse an uploaded document without persisting personal material to disk."""
    suffix = Path(filename).suffix.casefold()
    if suffix not in LOADERS:
        supported = ", ".join(sorted(LOADERS))
        raise UnsupportedDocumentError(
            f"Unsupported document type '{suffix}'. Supported: {supported}"
        )

    if suffix in {".md", ".markdown", ".txt"}:
        try:
            extracted = content.decode("utf-8")
        except UnicodeDecodeError as error:
            raise ValueError("Text uploads must use UTF-8 encoding.") from error
    elif suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as error:
            raise RuntimeError("Install document extras with: uv sync --extra documents") from error
        try:
            extracted = "\n\n".join(
                page.extract_text() or "" for page in PdfReader(BytesIO(content)).pages
            )
        except Exception as error:
            raise ValueError(f"Could not parse PDF upload: {filename}") from error
    else:
        try:
            from docx import Document
        except ImportError as error:
            raise RuntimeError("Install document extras with: uv sync --extra documents") from error
        try:
            document = Document(BytesIO(content))
            blocks = [
                paragraph.text.strip()
                for paragraph in document.paragraphs
                if paragraph.text.strip()
            ]
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        blocks.append(" | ".join(cells))
            extracted = "\n\n".join(blocks)
        except Exception as error:
            raise ValueError(f"Could not parse DOCX upload: {filename}") from error

    extracted = extracted.strip()
    if not extracted:
        raise ValueError(f"No extractable text found in: {filename}")
    return EvidenceDocument(
        source_id=_uploaded_source_id(filename, content),
        source_path=filename,
        content=extracted,
    )

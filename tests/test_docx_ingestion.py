from pathlib import Path

from docx import Document

from app.services.ingestion import load_document


def test_load_docx_extracts_paragraphs_and_tables(tmp_path: Path) -> None:
    source = tmp_path / "resume.docx"
    document = Document()
    document.add_paragraph("Built a Python API.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Skill"
    table.cell(0, 1).text = "FastAPI"
    document.save(source)

    loaded = load_document(source)

    assert "Built a Python API." in loaded.content
    assert "Skill | FastAPI" in loaded.content
